from api.app import mongodb
from api.app.transaction import transaction_bp
from api.app.forms.form_models import new_transaction_model
import plotly.express as px
import plotly.graph_objs as go
import pandas as pd
from flask import render_template, request, redirect, url_for, make_response
import json
import plotly
from plotly.graph_objs import *
import docx
import io
import matplotlib.pyplot as plt
from helpers.cf_analytics import update_cf, date_adj
import datetime as dt
from gridfs import GridFS
from helpers.function_hub import convert_pdf_to_text_sentences

db = mongodb['myDatabase']
coll = db['deals']

@transaction_bp.route("/transaction/dashboard", methods = ["GET", "POST"])
def render_dashboard():

    values = {}
    liquidity = {}
    ptf_count = 0

    red = 'rgb(255,0,0)'
    blue = 'rgb(0,0,255)'

    for deals in coll.find():

        ptf_count += 1
        values[deals['asset_name']] = deals['transaction_value']
        liquidity[deals['asset_name']] = deals['liquidity_type']

    deals_ser = pd.Series(values)
    liquidity_ser = pd.Series(liquidity)
    cf_data = update_cf().sort_index(ascending=True)

    clrs = [red if x < 0 else blue for x in cf_data]

    fig = px.pie(title = 'Asset overview', values = deals_ser, names=deals_ser.index.values)
    fig.update_layout({
        'paper_bgcolor' : 'rgba(0,0,0,0)',
        'plot_bgcolor' : 'rgba(0,0,0,0)'
    })

    fig = json.dumps(fig, cls=plotly.utils.PlotlyJSONEncoder)

    fig1 = px.pie(title = 'Liquidity overview', values = liquidity_ser, names = liquidity_ser.index.values)
    fig1 = json.dumps(fig1, cls = plotly.utils.PlotlyJSONEncoder)

    fig4 = go.Bar(x = cf_data.index.values,
                  y = cf_data.values,
                  marker=dict(color = clrs))

    fig4 = go.Figure(data = fig4)
    fig4 = json.dumps(fig4, cls = plotly.utils.PlotlyJSONEncoder)

    form = new_transaction_model()

    #fig4 = px.bar(cf_data, title = 'Cashflow')
    #fig4 = json.dumps(fig4, cls=plotly.utils.PlotlyJSONEncoder)

    return render_template('dashboard.html', form = form, fig = fig, fig1 = fig1, fig4 = fig4, table = [deals_ser.to_frame().to_html(classes = 'data')])

@transaction_bp.route("/render_form", methods = ["GET", "POST"])
def render_form():

    form = new_transaction_model()
    file = None
    #form['asset_name'].data = 'FundRisQ S.A.' #to be used when pre-filling the form (for example when using a divestment icon next to each asset to prefill form with asset name & divestment flag)

    if request.method == "POST":
        fund_data = form.data.copy()
        if fund_data['transaction_type'] == 'Divestment' and fund_data['transaction_type'] > 0:
            fund_data['transaction_amount'] = - fund_data['transaction_amount']
        fund_data['asset_id'] = coll.find_one(sort = [("asset_id", -1)])['asset_id'] + 1

        file = request.files.get('upload_memo')
        a = file.filename

        #file = fund_data['upload_memo'].read()
        #with open(fund_data['upload_memo'], 'rb') as doc:
        storage = GridFS(db, 'deals')
        storage.put(file, asset_id = fund_data['asset_id'], filename = a)

        pdf = convert_pdf_to_text_sentences(file)
        db['deals.files'].update_one({'asset_id': fund_data['asset_id']}, {"$set": {'pdf_text': pdf}})

        del fund_data['upload_memo']

        coll.insert_one(fund_data)

    return render_template('new_transaction.html', form = form)

@transaction_bp.route("/download", methods = ["GET", "POST"])
def download():

    storage = GridFS(db, 'deals')
    data = storage.get_last_version('schroder-gaia---bluetrend---de.pdf')
    response = make_response(data.read())
    extension = data.filename.split('.')[-1]
    response.headers['Content-Type'] = f'application/{extension}'
    response.headers['Content-Disposition'] = f'inline; filename={data.filename}'
    return response

@transaction_bp.route("/generate_report", methods = ["GET", "POST"])
def generate_report():

    values = {}
    liquidity = {}

    document = docx.Document()

    document.add_heading('Quarterly risk report', 0)
    document.add_heading('Subfund XYZ', 2)

    document.add_paragraph('This risk report provides you with an overview of the subfund"s exposure towards market, credit, liquidity and other risks. Please contact the risk department if you have further questions')

    document.add_paragraph()
    document.add_heading('Portfolio overview', 3)

    for deals in coll.find():

        values[deals['asset_name']] = deals['transaction_value']
        liquidity[deals['asset_name']] = deals['liquidity_type']

    deals_ser = pd.Series(values)
    liquidity_ser = pd.Series(liquidity)

    memfile = io.BytesIO()
    plt.plot(deals_ser)
    plt.savefig(memfile)

    #fig = px.pie(values = deals_ser, names=deals_ser.index.values)
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture(memfile)

    document.save('riskreport.docx')
    print('Report has been generated')

    return redirect(url_for('transaction_bp.render_dashboard'))

@transaction_bp.route("/portfolio/", methods = ["GET", "POST"])
def portfolio():

    dict = {}

    form = new_transaction_model()
    fund_data = form.data.copy()

    if fund_data['funds'] != "":
        _id = fund_data['funds']
        asset_id = _id
        return redirect(url_for("transaction_bp.overview", asset_id = asset_id))

    else:

        for k in coll.find():

            dict[k['asset_name']] = {
                'Asset Value' : k['transaction_value'],
                'Location' : k['asset_loc'],
                'Asset Leverage' : k['asset_leverage'],
                'Memo': 1,
                'Asset Id': k['asset_id']
            }

        df = pd.DataFrame(dict)
        df = df.reset_index()

        return render_template('portfolio.html', column_names = df.columns.values, row_data=df.values, zip=zip, len = len(df.columns.values))

@transaction_bp.route("/overview/<asset_id>", methods = ["GET"])
def overview(asset_id):

    dict = {}
    k = coll.find_one({'asset_id': int(asset_id)})

    dict[k['asset_name']] = {
            'Asset Value' : k['transaction_value'],
            'Location' : k['asset_loc'],
            'Asset Leverage' : k['asset_leverage'],
            'Memo': 1,
            'Asset Id': k['asset_id']
        }

    df = pd.DataFrame(dict)
    df = df.reset_index()

    return render_template('portfolio.html', column_names = df.columns.values, row_data=df.values, zip=zip, len = len(df.columns.values))

@transaction_bp.route("/home", methods = ["GET", "POST"])
def home():

    dict = {}
    x = 0
    funds = coll.distinct("fund_subfund_name")
    portfolio = {}
    master_df = pd.DataFrame()

    for f in funds:
        AuM = 0
        Leverage = 0

    for deals in coll.find():
        x = deals["transaction_value"]
        for deals in coll:
            if deals['fund_subfund_name'] == f:
                AuM += deals['transaction_value']
                Leverage += deals['asset_leverage']

        portfolio[f] = {'AuM': AuM, 'Leverage': Leverage}
        df = pd.DataFrame(data = portfolio.values(), index = portfolio.keys())
        #master_df.append(df)
        master_df = pd.concat([master_df, df])

    news_coll = db['news']



    return render_template("home.html", value = x)










